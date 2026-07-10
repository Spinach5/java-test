#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Java 实验一：Java 开发环境配置 实验报告 docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/workspace/24软件4班_曾秦伟2_第1次实验报告.docx"

doc = Document()

# 设置全局字体（宋体）
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
rpr = style.element.rPr
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:eastAsia'), '宋体')
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')

# 页面边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_title(text, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    rpr2 = run._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), '黑体')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rpr2.append(rf)
    return p


def add_heading(text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    rpr2 = run._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), '黑体')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rpr2.append(rf)
    return p


def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    rpr2 = run._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), '宋体')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rpr2.append(rf)
    return p


def add_code(text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    rpr2 = run._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), 'Consolas')
    rf.set(qn('w:ascii'), 'Consolas')
    rf.set(qn('w:hAnsi'), 'Consolas')
    rpr2.append(rf)
    # 浅灰背景
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr = p._element.get_or_add_pPr()
    pPr.append(shd)
    return p


# ==================== 封面信息 ====================
add_title("实验1：Java 开发环境配置", size=20)

# 表格：实验信息
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = [
    ("课程名称", "Java 程序设计", "实验名称", "Java 开发环境配置"),
    ("学院", "计算机学院", "专业班级", "24 软件 4 班"),
    ("姓名", "曾秦伟 2", "学号", "（请填写本人学号）"),
    ("指导教师", "（请填写）", "实验日期", "2026 年 7 月 10 日"),
]
for i, row in enumerate(info):
    cells = table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val
        for para in cells[j].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(11)
                run.bold = (j == 0 or j == 2)

doc.add_paragraph()

# ==================== 一、实验目的 ====================
add_heading("一、实验目的", level=1)
add_para("1. 掌握 JDK 的下载、安装与环境变量配置，熟悉 Java 开发环境的搭建流程。")
add_para("2. 熟悉常用 Java 集成开发环境（IntelliJ IDEA / Eclipse）的基本使用。")
add_para("3. 认识标准 Java 程序的基本结构，理解程序各组成部分的功能。")

# ==================== 二、实验内容 ====================
add_heading("二、实验内容", level=1)
add_para("1. 完成 JDK 的安装与系统环境变量配置，验证 Java 开发环境是否配置成功。")
add_para("2. 使用文本编辑器编写 HelloWorld.java 程序，通过命令行完成编译与运行。")
add_para("3. 选择一款 Java IDE（IntelliJ IDEA / Eclipse），创建项目并编写、编译、调试 HelloWorld 程序。")

# ==================== 三、实验环境 ====================
add_heading("三、实验环境", level=1)
add_para("操作系统：Windows 11（64 位）")
add_para("JDK 版本：JDK 17.0.10（长期支持版）")
add_para("开发工具：IntelliJ IDEA 2024.1 Community Edition / 命令行（PowerShell）")
add_para("辅助工具：Notepad++（文本编辑器）")

# ==================== 四、实验步骤 ====================
add_heading("四、实验步骤与过程", level=1)

# 4.1
add_heading("4.1 JDK 的下载与安装", level=2)
add_para("（1）打开浏览器，访问 Oracle 官方下载页面：https://www.oracle.com/java/technologies/downloads/，选择 Windows 操作系统对应的 JDK 17 安装包 jdk-17_windows-x64_bin.exe 进行下载。", indent=False)
add_para("（2）下载完成后，双击安装包，弹出安装向导。点击“下一步”进入安装选项设置，自定义安装目录为：", indent=False)
add_code("D:\\Program Files\\Java\\jdk-17")
add_para("（3）安装程序同时会提示安装 JRE（Java 运行时环境），可继续使用默认路径：", indent=False)
add_code("D:\\Program Files\\Java\\jre-17")
add_para("（4）等待安装进度条完成，点击“关闭”结束安装过程。", indent=False)
add_para("（5）进入 JDK 安装目录的 bin 文件夹（D:\\Program Files\\Java\\jdk-17\\bin），可以看到 java.exe、javac.exe、jar.exe 等可执行文件，说明 JDK 已成功安装。", indent=False)

# 4.2
add_heading("4.2 环境变量配置", level=2)
add_para("右键点击“此电脑” → “属性” → “高级系统设置” → “环境变量”，分别在用户变量或系统变量中新增以下三项：", indent=False)
add_para("① 新建 JAVA_HOME，变量值为 JDK 的安装路径：", indent=False)
add_code("JAVA_HOME = D:\\Program Files\\Java\\jdk-17")
add_para("② 新建 CLASSPATH，变量值为 JDK 类库搜索路径：", indent=False)
add_code("CLASSPATH = .;%JAVA_HOME%\\lib\\dt.jar;%JAVA_HOME%\\lib\\tools.jar")
add_para("③ 编辑 Path 变量，在其值的最前面添加 JDK 的 bin 目录：", indent=False)
add_code("%JAVA_HOME%\\bin")
add_para("（注意：从 JDK 9 开始，dt.jar 和 tools.jar 已被移除，实际配置时 CLASSPATH 可简写为“.”即可。）", indent=False)
add_para("设置完毕后，依次点击“确定”保存所有对话框，使环境变量生效。", indent=False)

# 4.3
add_heading("4.3 验证 Java 开发环境", level=2)
add_para("使用快捷键 Win + R 打开“运行”窗口，输入 cmd 回车打开命令行窗口，依次输入以下三条命令进行验证：", indent=False)
add_code("java -version")
add_code("javac -version")
add_code("echo %JAVA_HOME%")
add_para("若环境配置成功，命令行将分别显示类似如下的输出：", indent=False)
add_code("java version \"17.0.10\" 2024-01-16\nJava(TM) SE Runtime Environment (build 17.0.10+8)\nJava HotSpot(TM) 64-Bit Server VM (build 17.0.10+8, mixed mode, sharing)")
add_code("javac 17.0.10")
add_code("D:\\Program Files\\Java\\jdk-17")
add_para("若三条命令均能正确返回信息，则说明 JDK 安装与环境变量配置成功。", indent=False)
add_para("（附：此处应附上命令行运行结果截图，截图文件命名为 验证截图.png 并打包提交。）", indent=False)

# 4.4
add_heading("4.4 命令行方式编写并运行 HelloWorld 程序", level=2)
add_para("（1）在 D 盘根目录下新建一个工作目录 D:\\JavaWork，并在其中新建一个文本文档，文件名为 HelloWorld.java。", indent=False)
add_para("（2）使用 Notepad++ 打开该文件，输入如下 Java 源代码：", indent=False)
add_code("public class HelloWorld {\n    public static void main(String[] args) {\n        System.out.println(\"Hello, World! 欢迎来到 Java 世界！\");\n        System.out.println(\"姓名：曾秦伟 2    班级：24软件4班\");\n        System.out.println(\"Java 是一门面向对象的、跨平台的程序设计语言。\");\n    }\n}")
add_para("（3）保存文件后，打开命令行窗口，使用 cd 命令切换到源文件所在目录：", indent=False)
add_code("cd D:\\JavaWork")
add_para("（4）执行 javac 命令编译源文件，生成字节码文件：", indent=False)
add_code("javac HelloWorld.java")
add_para("编译成功后，目录中会生成 HelloWorld.class 字节码文件。", indent=False)
add_para("（5）执行 java 命令运行程序，虚拟机加载并执行字节码：", indent=False)
add_code("java HelloWorld")
add_para("（6）控制台将输出如下内容：", indent=False)
add_code("Hello, World! 欢迎来到 Java 世界！\n姓名：曾秦伟 2    班级：24软件4班\nJava 是一门面向对象的、跨平台的程序设计语言。")
add_para("（附：此处应附上编译及运行结果截图，文件命名为 命令行运行截图.png。）", indent=False)

# 4.5
add_heading("4.5 使用 IntelliJ IDEA 创建项目并调试", level=2)
add_para("（1）从 JetBrains 官网下载 IntelliJ IDEA 2024.1 Community 社区版，按向导完成安装。", indent=False)
add_para("（2）启动 IDEA，在欢迎界面点击 “New Project”，选择左侧的 “Java”，确认 Project SDK 已自动识别为 JDK 17（若未识别可手动指定），点击 “Next”。", indent=False)
add_para("（3）勾选 “Create project from template” 中的 “Command Line App” 模板（也可不勾选自行创建），项目命名为 HelloWorldDemo，存放位置选择 D:\\JavaWork\\IDEAProject，点击 “Finish”。", indent=False)
add_para("（4）IDE 自动生成 src 文件夹与 Main.java，删除 Main.java，右键点击 src 目录 → New → Java Class，类名填写 HelloWorld，回车确认。", indent=False)
add_para("（5）在编辑窗口中输入 4.4 步骤中的 HelloWorld 程序代码，IDEA 会自动保存。", indent=False)
add_para("（6）设置断点：将光标移至 System.out.println 语句行，点击行号左侧灰色区域，出现红色圆点即为断点。", indent=False)
add_para("（7）右键点击编辑区 → “Debug 'HelloWorld.main()'”，程序进入调试模式，停在断点处。可使用 F8（Step Over）逐行执行、F9（Resume Program）继续执行等。", indent=False)
add_para("（8）调试完成后，Run 工具栏点击绿色三角“Run”按钮运行程序，下方 Run 窗口显示输出结果，验证程序功能。", indent=False)
add_para("（附：此处应附 IDEA 项目结构、运行结果及调试窗口截图，分别命名为 IDEA项目结构.png、IDEA运行结果.png、IDEA调试截图.png。）", indent=False)

# ==================== 五、程序结构解析 ====================
add_heading("五、HelloWorld 程序结构逐行解析", level=1)

add_para("下面对 HelloWorld.java 中的每一行代码进行详细分析：", indent=False)
add_code("public class HelloWorld {")
add_para("【第 1 行】class 是 Java 中定义类的关键字；HelloWorld 是类名，类名必须与源文件名 HelloWorld.java 完全保持一致，且首字母大写；public 是访问修饰符，表示该类可以被任意位置的类访问；左花括号 { 表示类体的开始。", indent=False)

add_code("    public static void main(String[] args) {")
add_para("【第 2 行】定义 Java 程序的入口方法 main：", indent=False)
add_para("    • public：访问修饰符，JVM 必须能从外部调用此方法。", indent=False)
add_para("    • static：静态修饰符，表明该方法属于类本身，可通过类名直接调用，无需创建对象。", indent=False)
add_para("    • void：返回值类型，表示该方法没有返回值。", indent=False)
add_para("    • main：方法名，是 Java 应用程序的固定入口方法名，不可更改。", indent=False)
add_para("    • String[] args：字符串数组参数，用于接收命令行传入的参数。", indent=False)
add_para("    • 左花括号 { 表示方法体的开始。", indent=False)

add_code("        System.out.println(\"Hello, World! 欢迎来到 Java 世界！\");")
add_para("【第 3 行】System 是 java.lang 包中提供的标准类；out 是 System 类的静态成员，代表标准输出流（控制台）；println 是 PrintStream 类的方法，作用是向控制台输出字符串并换行；字符串内容为程序的问候信息。整行语句以分号 ; 结束，Java 中每条语句必须以分号结尾。", indent=False)

add_code("        System.out.println(\"姓名：曾秦伟 2    班级：24软件4班\");")
add_para("【第 4 行】与第 3 行功能相同，输出包含本人姓名与班级信息，用于标识输出者。", indent=False)

add_code("        System.out.println(\"Java 是一门面向对象的、跨平台的程序设计语言。\");")
add_para("【第 5 行】同上，输出对 Java 语言的简要说明，进一步体现程序的输出功能。", indent=False)

add_code("    }")
add_para("【第 6 行】右花括号 } 表示 main 方法体的结束，与第 2 行的左花括号配对。", indent=False)

add_code("}")
add_para("【第 7 行】右花括号 } 表示 HelloWorld 类体的结束，与第 1 行的左花括号配对。", indent=False)

add_para("综上，HelloWorld 程序主要由 4 个核心部分组成：", indent=False)
add_para("① 类的定义（class HelloWorld）。", indent=False)
add_para("② 入口方法 main(String[] args)。", indent=False)
add_para("③ 调用标准输出语句 System.out.println(...) 输出信息。", indent=False)
add_para("④ 类与方法的花括号 { } 配对表示作用域范围。", indent=False)

# ==================== 六、实验结果 ====================
add_heading("六、实验结果", level=1)
add_para("1. 成功安装 JDK 17，并在 Windows 11 系统中正确配置 JAVA_HOME、CLASSPATH 与 Path 环境变量，命令行执行 java -version、javac -version、echo %JAVA_HOME% 均可正确返回版本信息，开发环境配置成功。", indent=False)
add_para("2. 成功使用 Notepad++ 编写 HelloWorld.java，并通过 javac 与 java 命令完成编译与运行，控制台正确输出三行问候信息。", indent=False)
add_para("3. 成功使用 IntelliJ IDEA 2024.1 创建 Java 项目、编写 HelloWorld 程序，并使用调试功能进行断点调试，运行结果与命令行结果一致。", indent=False)

# ==================== 七、问题与解决 ====================
add_heading("七、实验过程中遇到的问题、原因与解决方法", level=1)

add_para("【问题 1】配置完环境变量后，命令行输入 java -version 仍然提示 “'java' 不是内部或外部命令”。", indent=False)
add_para("原因：修改环境变量后没有重新打开命令行窗口，旧 cmd 进程的环境变量缓存未刷新；或在 Path 中使用了错误的变量引用方式。", indent=False)
add_para("解决：关闭当前所有命令行窗口，重新打开 cmd，再次执行 java -version。同时确认 Path 中填入的是 %JAVA_HOME%\\bin 而不是字面值，并且百分号、斜杠方向都正确。", indent=False)

add_para("【问题 2】执行 javac HelloWorld.java 编译时报错 “错误: 找不到或无法加载主类 HelloWorld”。", indent=False)
add_para("原因：执行 java 命令时，当前路径下没有对应的 .class 文件，或类名与文件名不一致（如源文件命名为 Helloworld.java，但类名为 HelloWorld）。", indent=False)
add_para("解决：使用 dir 命令检查当前目录是否生成了 HelloWorld.class 字节码文件；确认源文件名、类名、执行的类名三者完全一致，且首字母大小写相同。", indent=False)

add_para("【问题 3】使用 IntelliJ IDEA 创建项目时，Project SDK 显示为 “No SDK”。", indent=False)
add_para("原因：IDEA 启动时未自动找到已安装的 JDK。", indent=False)
add_para("解决：点击右侧 “…” 按钮，弹出 Choose SDK 窗口，选择 “Add JDK”，定位到 D:\\Program Files\\Java\\jdk-17 目录并确认，SDK 即被正确识别。", indent=False)

add_para("【问题 4】运行 HelloWorld.java 时控制台输出中文乱码。", indent=False)
add_para("原因：源文件编码与控制台/IDEA 默认编码不一致（如源文件为 GBK，控制台为 UTF-8），导致中文字符无法正确显示。", indent=False)
add_para("解决：在保存源文件时使用 UTF-8 编码；若使用命令行方式，先执行 chcp 65001 切换到 UTF-8 代码页；若使用 IDEA，则在 File → Settings → Editor → File Encodings 中将 Global Encoding、Project Encoding、Default encoding for properties files 都设置为 UTF-8 即可。", indent=False)

# ==================== 八、实验总结 ====================
add_heading("八、实验总结", level=1)
add_para("通过本次实验，我系统地学习了 Java 开发环境的搭建流程，主要包括以下收获：", indent=False)
add_para("1. 掌握了 JDK 17 的下载、安装步骤，理解了 JDK、JRE、JVM 三者的关系：JDK 是开发工具包，JRE 是运行时环境，JVM 是 Java 虚拟机，JDK 中包含 JRE，JRE 中又包含 JVM。", indent=False)
add_para("2. 熟练配置了 JAVA_HOME、CLASSPATH、Path 三个核心环境变量，并通过 java -version、javac -version 等命令验证了环境配置的正确性。", indent=False)
add_para("3. 掌握了通过命令行 javac 与 java 命令进行 Java 程序的编译与运行，理解了源文件（.java）经编译生成字节码文件（.class），再由 JVM 解释执行的完整过程。", indent=False)
add_para("4. 学会了使用 IntelliJ IDEA 创建 Java 项目、编写代码、设置断点与调试程序，初步体验了现代化 IDE 带来的高效开发体验。", indent=False)
add_para("5. 通过对 HelloWorld 程序的逐行解析，加深了对 Java 程序结构（类、方法、语句、注释、花括号作用域等）的认识，为后续深入学习 Java 面向对象编程打下了坚实基础。", indent=False)
add_para("本次实验过程中也遇到了一些典型问题（如环境变量不生效、类名文件名不一致、中文乱码等），通过查阅资料、独立排查都得到了有效解决，提升了自主学习与动手实践能力。", indent=False)

# ==================== 参考文献 ====================
add_heading("参考文献", level=1)
add_para("[1] 黑马程序员. Java 基础入门（第 3 版）[M]. 北京：清华大学出版社, 2023.", indent=False)
add_para("[2] Oracle. The Java Tutorials – Getting Started[EB/OL]. https://docs.oracle.com/javase/tutorial/getStarted/index.html, 2026.", indent=False)
add_para("[3] JetBrains. IntelliJ IDEA Documentation[EB/OL]. https://www.jetbrains.com/help/idea/getting-started.html, 2026.", indent=False)

# 保存
doc.save(OUTPUT)
print(f"实验报告已成功生成: {OUTPUT}")
